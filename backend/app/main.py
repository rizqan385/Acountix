# pyright: reportAttributeAccessIssue=false
import json
import os
import re
import uuid
import shutil
import time
from pathlib import Path
from typing import Any

from fastapi import FastAPI, HTTPException, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from fastapi.staticfiles import StaticFiles
from google.genai import types

# Import modul buatan kita sendiri dari file sebelah
from app.config import gemini_client, groq_client, supabase
from app.schemas import (
    SoalRequest, ChatRequest, ChatResponse,
    HasilParserAkuntansi, EntriJurnal, ResponAIPintar
)

app = FastAPI(title="Acountix AI Engine - Clean Architecture", version="3.0.0")

# CORS biar frontend bisa ngobrol sama backend
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Folder buat nyimpen file upload & generated Excel
UPLOAD_DIR = Path("uploads")
UPLOAD_DIR.mkdir(exist_ok=True)
EXCEL_DIR = Path("generated_excel")
EXCEL_DIR.mkdir(exist_ok=True)

# Serve static files (generated Excel)
app.mount("/files", StaticFiles(directory=str(EXCEL_DIR)), name="files")

# ===== HELPER: Extract text dari berbagai format file =====
def extract_text_from_file(file_path: str, filename: str) -> str:
    """Extract teks dari file. Support: .doc, .docx, .txt, .pdf"""
    ext = Path(filename).suffix.lower()

    if ext == '.txt':
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()

    elif ext == '.docx':
        try:
            import docx
            doc = docx.Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            # Juga extract tabel
            for table in doc.tables:
                for row in table.rows:
                    row_text = '\t'.join(cell.text.strip() for cell in row.cells)
                    paragraphs.append(row_text)
            return '\n'.join(paragraphs)
        except Exception as e:
            raise ValueError(f"Gagal baca file .docx: {str(e)}")

    elif ext == '.doc':
        try:
            import olefile
            ole = olefile.OleFileIO(file_path)
            if ole.exists('WordDocument'):
                data = ole.openstream('WordDocument').read()
                text = data.decode('latin-1', errors='ignore')
                # Bersihin karakter binary, pertahanin karakter cetak
                clean = re.sub(r'[^\x20-\x7E\n\r\tÀ-ÿ]', ' ', text)
                clean = re.sub(r' {3,}', ' ', clean)
                # Filter: cari awal konten bermakna, buang header biner OLE
                # Split pakai \r dan \n supaya biner dan teks yang menyatu bisa terpisah
                lines = re.split(r'[\r\n]+', clean)
                meaningful_lines = []
                content_started = False
                # Kata kunci penanda awal konten dokumen Indonesia
                content_markers = ['perusahaan', 'soal', 'latihan', 'neraca', 'jurnal', 
                                   'transaksi', 'pembelian', 'penjualan', 'januari', 'februari',
                                   'maret', 'april', 'mei', 'juni', 'juli', 'agustus',
                                   'september', 'oktober', 'november', 'desember',
                                   'manufaktur', 'akuntansi']
                for line in lines:
                    stripped = line.strip()
                    if len(stripped) < 5:
                        continue
                    if not content_started:
                        lower_line = stripped.lower()
                        if any(marker in lower_line for marker in content_markers):
                            content_started = True
                    if content_started:
                        # Dual check: rasio ASCII tinggi DAN mengandung kata nyata (3+ huruf)
                        ascii_normal = sum(1 for c in stripped if c.isascii() and (c.isalnum() or c in ' .,;:!?@#%&()/-\t\'\"'))
                        ratio = ascii_normal / len(stripped)
                        real_words = re.findall(r'[A-Za-z]{3,}', stripped)
                        if ratio > 0.7 and len(real_words) >= 2:
                            meaningful_lines.append(stripped)
                return '\n'.join(meaningful_lines)
            else:
                raise ValueError("File .doc tidak mengandung data teks.")
        except ValueError:
            raise
        except Exception as e:
            raise ValueError(f"Gagal baca file .doc: {str(e)}")

    elif ext == '.pdf':
        # Coba pake pdfplumber kalau ada, kalau ga fallback ke pembacaan sederhana
        try:
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                text = ''
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + '\n'
                return text
        except ImportError:
            raise ValueError("Library pdfplumber belum terinstall. Install dulu: pip install pdfplumber")
        except Exception as e:
            raise ValueError(f"Gagal baca file PDF: {str(e)}")

    else:
        # Fallback: coba baca sebagai text
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
        except Exception:
            raise ValueError(f"Format file '{ext}' belum didukung. Coba pake .doc, .docx, .txt, atau .pdf ya!")


# ===== HELPER: Generate Excel dari data jurnal =====
def generate_excel(data: dict, user_name: str) -> str:
    """Generate file Excel dari data jurnal. Return filename."""
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    wb = Workbook()

    # ===== SHEET 1: Jurnal Umum (Terisi) =====
    ws: Any = wb.active
    ws.title = "Jurnal Umum"

    # Styling
    header_font = Font(name='Calibri', bold=True, size=12, color='FFFFFF')
    header_fill = PatternFill(start_color='BE185D', end_color='BE185D', fill_type='solid')
    title_font = Font(name='Calibri', bold=True, size=14, color='BE185D')
    data_font = Font(name='Calibri', size=11)
    money_font = Font(name='Calibri', size=11, color='1F7A1F')
    border = Border(
        left=Side(style='thin', color='D4D4D4'),
        right=Side(style='thin', color='D4D4D4'),
        top=Side(style='thin', color='D4D4D4'),
        bottom=Side(style='thin', color='D4D4D4'),
    )
    center = Alignment(horizontal='center', vertical='center')
    wrap = Alignment(horizontal='left', vertical='center', wrap_text=True)

    # Title
    ws.merge_cells('A1:E1')
    ws['A1'] = f"JURNAL UMUM - {data.get('nama_perusahaan', 'N/A')}"
    ws['A1'].font = title_font  # type: ignore
    ws['A1'].alignment = center  # type: ignore

    ws.merge_cells('A2:E2')
    ws['A2'] = f"Periode: {data.get('periode', 'N/A')}"
    ws['A2'].font = Font(name='Calibri', size=11, italic=True, color='666666')  # type: ignore
    ws['A2'].alignment = center  # type: ignore

    # Header row
    headers = ['Tanggal', 'Akun Debit', 'Nominal Debit', 'Akun Kredit', 'Nominal Kredit']
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=4, column=col, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = center
        cell.border = border

    # Data rows
    for i, jurnal in enumerate(data.get('daftar_jurnal', []), 5):
        ws.cell(row=i, column=1, value=jurnal.get('tanggal', '')).font = data_font
        ws.cell(row=i, column=2, value=jurnal.get('akun_debit', '')).font = data_font
        ws.cell(row=i, column=3, value=jurnal.get('nominal_debit', 0)).font = money_font
        ws.cell(row=i, column=3).number_format = '#,##0'
        ws.cell(row=i, column=4, value=jurnal.get('akun_kredit', '')).font = data_font
        ws.cell(row=i, column=5, value=jurnal.get('nominal_kredit', 0)).font = money_font
        ws.cell(row=i, column=5).number_format = '#,##0'
        for col in range(1, 6):
            ws.cell(row=i, column=col).border = border
            ws.cell(row=i, column=col).alignment = wrap

    # Total row
    total_row = 5 + len(data.get('daftar_jurnal', []))
    total_debit = sum(j.get('nominal_debit', 0) for j in data.get('daftar_jurnal', []))
    total_kredit = sum(j.get('nominal_kredit', 0) for j in data.get('daftar_jurnal', []))

    ws.cell(row=total_row, column=1, value='TOTAL').font = Font(name='Calibri', bold=True, size=11)
    ws.cell(row=total_row, column=3, value=total_debit).font = Font(name='Calibri', bold=True, size=11, color='1F7A1F')
    ws.cell(row=total_row, column=3).number_format = '#,##0'
    ws.cell(row=total_row, column=5, value=total_kredit).font = Font(name='Calibri', bold=True, size=11, color='1F7A1F')
    ws.cell(row=total_row, column=5).number_format = '#,##0'

    for col in range(1, 6):
        ws.cell(row=total_row, column=col).border = Border(
            top=Side(style='double', color='BE185D'),
            bottom=Side(style='double', color='BE185D'),
        )

    # Column widths
    ws.column_dimensions['A'].width = 15
    ws.column_dimensions['B'].width = 30
    ws.column_dimensions['C'].width = 20
    ws.column_dimensions['D'].width = 30
    ws.column_dimensions['E'].width = 20

    # ===== SHEET 2: Template Kosong (buat user isi sendiri) =====
    ws2: Any = wb.create_sheet("Template Jurnal (Kosong)")

    ws2.merge_cells('A1:E1')
    ws2['A1'] = f"TEMPLATE JURNAL - {data.get('nama_perusahaan', 'Isi Nama Perusahaan')}"
    ws2['A1'].font = title_font  # type: ignore
    ws2['A1'].alignment = center  # type: ignore

    ws2.merge_cells('A2:E2')
    ws2['A2'] = "Silakan isi jurnal di bawah ini ✏️"
    ws2['A2'].font = Font(name='Calibri', size=11, italic=True, color='666666')  # type: ignore
    ws2['A2'].alignment = center  # type: ignore

    for col, header in enumerate(headers, 1):
        cell = ws2.cell(row=4, column=col, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = center
        cell.border = border

    # Kosongkan baris buat diisi (20 baris template)
    for row in range(5, 25):
        for col in range(1, 6):
            cell = ws2.cell(row=row, column=col, value='')
            cell.border = border
            cell.font = data_font
        # Format kolom nominal sebagai angka
        ws2.cell(row=row, column=3).number_format = '#,##0'
        ws2.cell(row=row, column=5).number_format = '#,##0'

    ws2.column_dimensions['A'].width = 15
    ws2.column_dimensions['B'].width = 30
    ws2.column_dimensions['C'].width = 20
    ws2.column_dimensions['D'].width = 30
    ws2.column_dimensions['E'].width = 20

    # ===== SHEET 3: Neraca Saldo (opsional) =====
    if data.get('neraca_saldo'):
        ws3: Any = wb.create_sheet("Neraca Saldo")
        ws3.merge_cells('A1:C1')
        ws3['A1'] = f"NERACA SALDO - {data.get('nama_perusahaan', 'N/A')}"
        ws3['A1'].font = title_font  # type: ignore
        ws3['A1'].alignment = center  # type: ignore
        
        ws3.merge_cells('A2:C2')
        ws3['A2'] = f"Periode: {data.get('periode', 'N/A')}"
        ws3['A2'].font = Font(name='Calibri', size=11, italic=True, color='666666')  # type: ignore
        ws3['A2'].alignment = center  # type: ignore
        
        headers_ns = ['Nama Akun', 'Debit', 'Kredit']
        for col, header in enumerate(headers_ns, 1):
            cell = ws3.cell(row=4, column=col, value=header)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = center
            cell.border = border
            
        for i, ns in enumerate(data['neraca_saldo'], 5):
            ws3.cell(row=i, column=1, value=ns.get('akun', '')).font = data_font
            ws3.cell(row=i, column=2, value=ns.get('debit', 0)).font = money_font
            ws3.cell(row=i, column=2).number_format = '#,##0'
            ws3.cell(row=i, column=3, value=ns.get('kredit', 0)).font = money_font
            ws3.cell(row=i, column=3).number_format = '#,##0'
            for col in range(1, 4):
                ws3.cell(row=i, column=col).border = border
                ws3.cell(row=i, column=col).alignment = wrap
                
        ws3.column_dimensions['A'].width = 30
        ws3.column_dimensions['B'].width = 20
        ws3.column_dimensions['C'].width = 20

    # ===== SHEET 4: Laporan HPP (opsional) =====
    if data.get('laporan_harga_pokok_produksi'):
        ws4: Any = wb.create_sheet("Laporan HPP")
        ws4.merge_cells('A1:B1')
        ws4['A1'] = f"LAPORAN HARGA POKOK PRODUKSI - {data.get('nama_perusahaan', 'N/A')}"
        ws4['A1'].font = title_font  # type: ignore
        ws4['A1'].alignment = center  # type: ignore
        
        ws4.merge_cells('A2:B2')
        ws4['A2'] = f"Periode: {data.get('periode', 'N/A')}"
        ws4['A2'].font = Font(name='Calibri', size=11, italic=True, color='666666')  # type: ignore
        ws4['A2'].alignment = center  # type: ignore
        
        for i, lap in enumerate(data['laporan_harga_pokok_produksi'], 4):
            ws4.cell(row=i, column=1, value=lap.get('keterangan', '')).font = data_font
            ws4.cell(row=i, column=2, value=lap.get('nominal', 0)).font = money_font
            ws4.cell(row=i, column=2).number_format = '#,##0'
            for col in range(1, 3):
                ws4.cell(row=i, column=col).border = border
                ws4.cell(row=i, column=col).alignment = wrap
                
        ws4.column_dimensions['A'].width = 40
        ws4.column_dimensions['B'].width = 20

    # ===== SHEET 5: Laporan HPPenjualan (opsional) =====
    if data.get('laporan_harga_pokok_penjualan'):
        ws5: Any = wb.create_sheet("Laporan HPPenjualan")
        ws5.merge_cells('A1:B1')
        ws5['A1'] = f"LAPORAN HARGA POKOK PENJUALAN - {data.get('nama_perusahaan', 'N/A')}"
        ws5['A1'].font = title_font  # type: ignore
        ws5['A1'].alignment = center  # type: ignore
        
        ws5.merge_cells('A2:B2')
        ws5['A2'] = f"Periode: {data.get('periode', 'N/A')}"
        ws5['A2'].font = Font(name='Calibri', size=11, italic=True, color='666666')  # type: ignore
        ws5['A2'].alignment = center  # type: ignore
        
        for i, lap in enumerate(data['laporan_harga_pokok_penjualan'], 4):
            ws5.cell(row=i, column=1, value=lap.get('keterangan', '')).font = data_font
            ws5.cell(row=i, column=2, value=lap.get('nominal', 0)).font = money_font
            ws5.cell(row=i, column=2).number_format = '#,##0'
            for col in range(1, 3):
                ws5.cell(row=i, column=col).border = border
                ws5.cell(row=i, column=col).alignment = wrap
                
        ws5.column_dimensions['A'].width = 40
        ws5.column_dimensions['B'].width = 20

    # Save
    safe_name = re.sub(r'[^\w\-]', '_', data.get('nama_perusahaan', 'jurnal'))
    filename = f"jurnal_{safe_name}_{uuid.uuid4().hex[:8]}.xlsx"
    filepath = EXCEL_DIR / filename
    wb.save(str(filepath))

    return filename


# ===== HELPER: Gemini API Retry Wrapper =====
def generate_content_with_retry(model, contents, config=None):
    """Bungkus call Gemini API dengan retry + auto-fallback ke Groq."""
    max_retries = 3
    base_delay = 2
    for attempt in range(max_retries):
        try:
            return gemini_client.models.generate_content(
                model=model,
                contents=contents,
                config=config
            )
        except Exception as e:
            err_msg = str(e).lower()
            is_rate_limit = any(k in err_msg for k in ["503", "429", "quota", "unavailable", "resource_exhausted"])
            if is_rate_limit:
                # Kalau punya Groq fallback, nggak usah nunggu lama-lama, langsung lempar error aja biar pindah ke Groq
                if groq_client:
                    print(f"[WARNING] Gemini kena limit. Skip waiting, langsung fallback ke Groq.")
                    raise e
                    
                if attempt < max_retries - 1:
                    sleep_time = base_delay * (2 ** attempt)
                    match = re.search(r'retry in ([\d\.]+)s', err_msg)
                    if match:
                        sleep_time = float(match.group(1)) + 1.0
                    print(f"[WARNING] Gemini kena limit. Retry in {sleep_time:.1f}s (Attempt {attempt+1}/{max_retries})")
                    time.sleep(sleep_time)
                    continue
            # Kalau udah habis retry atau error lain, raise supaya bisa di-catch oleh caller
            raise e


# ===== HELPER: Groq Fallback =====
def _call_groq_fallback(prompt: str) -> str:
    """Panggil Groq API sebagai fallback. Return raw JSON string."""
    if not groq_client:
        raise ValueError("Groq client tidak tersedia. Cek GROQ_API_KEY di .env")
    
    print("[INFO] 🔄 Switching ke Groq API sebagai fallback...")
    
    # Bungkus prompt dengan instruksi JSON format
    system_msg = """Kamu adalah Acountix, asisten AI akuntansi yang pintar. 
Kamu HARUS membalas dalam format JSON VALID (tanpa markdown code block, tanpa backtick).
JSON harus mengikuti skema ini:
{
  "is_akuntansi": boolean,
  "pesan_balasan": "string",
  "data_akuntansi": null atau {
    "nama_perusahaan": "string",
    "periode": "string",
    "daftar_jurnal": [{"tanggal": "string", "akun_debit": "string", "nominal_debit": number, "akun_kredit": "string", "nominal_kredit": number}],
    "neraca_saldo": null atau [{"akun": "string", "debit": number, "kredit": number}],
    "laporan_harga_pokok_produksi": null atau [{"keterangan": "string", "nominal": number}],
    "laporan_harga_pokok_penjualan": null atau [{"keterangan": "string", "nominal": number}]
  }
}
Jangan tambahkan teks apapun di luar JSON. Langsung mulai dengan { dan akhiri dengan }."""
    
    chat_completion = groq_client.chat.completions.create(
        messages=[
            {"role": "system", "content": system_msg},
            {"role": "user", "content": prompt}
        ],
        model="llama-3.3-70b-versatile",
        temperature=0.1,
        max_tokens=8000,
    )
    
    result = chat_completion.choices[0].message.content
    if not result:
        raise ValueError("Groq tidak mengembalikan respons.")
    
    # Bersihkan jika ada markdown code block wrapper
    result = result.strip()
    if result.startswith("```"):
        result = re.sub(r'^```(?:json)?\s*', '', result)
        result = re.sub(r'\s*```$', '', result)
    
    return result.strip()


# ===== HELPER: Proses AI Pintar (Unified) =====
def _process_ai_pintar(text_input: str, user_name: str, from_file: str = "") -> ChatResponse:
    prompt_instruksi = f'''
    Lu adalah Acountix, asisten AI akuntansi manufaktur yang super friendly, lucu, dan kece. 
    Tugas lu adalah membaca input dari user (bisa berupa chat biasa atau file berisi soal akuntansi)
    dan memberikan respons yang sesuai dalam format JSON terstruktur.

    Aturan:
    1. Tentukan apakah input berisi SOAL/TUGAS AKUNTANSI (ada transaksi, pembelian, nominal, dsb) atau KONTEN/CHAT BIASA.
    2. Jika KONTEN BIASA:
       - Set `is_akuntansi` = false
       - Isi `pesan_balasan` dengan obrolan asyik, penjelasan materi, atau sapaan ramah. Panggil user "{user_name}".
       - `data_akuntansi` biarkan kosong (null).
    3. Jika SOAL AKUNTANSI:
       - Set `is_akuntansi` = true
       - Isi `pesan_balasan` dengan pesan semangat atau sapaan singkat yang seru.
       - Isi `data_akuntansi` dengan parsing jurnal lengkap:
         - Jurnal Umum (HARUS balance, akrual ke 'Utang Beban', COA standar)
         - Neraca Saldo (jika ada data)
         - Laporan HPP & HPPenjualan (jika ada data)
         - Jika ada informasi nama perusahaan dan periode, masukkan ke data_akuntansi.
         - Kalau nama perusahaan ga disebut, pake "Perusahaan Tidak Diketahui".
         - Kalau periode ga disebut, pake bulan dan tahun sekarang.

    Input dari {user_name} / Isi File:
    {text_input}
    '''

    # Coba Gemini dulu, fallback ke Groq kalau kena rate limit
    try:
        response = generate_content_with_retry(
            model='gemini-2.5-flash',
            contents=prompt_instruksi,
            config=types.GenerateContentConfig(
                response_mime_type="application/json",
                response_schema=ResponAIPintar,
                temperature=0.1
            ),
        )
        if not response or not response.text:
            raise ValueError("Gemini tidak merespons.")
        ai_result = json.loads(response.text)
    except Exception as gemini_err:
        print(f"[WARNING] Gemini gagal: {str(gemini_err)[:80]}")
        if groq_client:
            raw_json = _call_groq_fallback(prompt_instruksi)
            ai_result = json.loads(raw_json)
            print("[INFO] ✅ Groq berhasil memproses request!")
        else:
            raise gemini_err
    
    if not ai_result.get("is_akuntansi"):
        return ChatResponse(
            type="chat",
            message=ai_result.get("pesan_balasan", "Hmm, aku nggak ngerti nih maksudnya apa.").strip(),
            data=None
        )
    
    # Kalau soal akuntansi
    data_hasil_ai = ai_result.get("data_akuntansi")
    if not data_hasil_ai or not data_hasil_ai.get('daftar_jurnal'):
        return ChatResponse(
            type="chat",
            message=ai_result.get("pesan_balasan", "Maaf banget, aku nggak nemu transaksi jurnal yang valid di dalamnya 😭"),
            data=None
        )

    # Validasi balance
    total_debit = sum(j.get('nominal_debit', 0) for j in data_hasil_ai.get('daftar_jurnal', []))
    total_kredit = sum(j.get('nominal_kredit', 0) for j in data_hasil_ai.get('daftar_jurnal', []))

    if total_debit != total_kredit:
        selisih = abs(total_debit - total_kredit)
        status_balance = f"⚠️ WARNING: Jurnal TIDAK BALANCE (Selisih: Rp {selisih:,.0f}). Harap cek Excel-nya!"
        # Tetap simpan ke database atau bisa di-skip, di sini kita tetap simpan agar user bisa revisi
    else:
        status_balance = "✅ Status: Jurnal BALANCE & tersimpan di database!"

    # Simpan ke Supabase
    for jurnal in data_hasil_ai['daftar_jurnal']:
        data_insert = {
            "nama_perusahaan": data_hasil_ai.get('nama_perusahaan', ''),
            "periode": data_hasil_ai.get('periode', ''),
            "tanggal": jurnal.get('tanggal', ''),
            "akun_debit": jurnal.get('akun_debit', ''),
            "nominal_debit": jurnal.get('nominal_debit', 0),
            "akun_kredit": jurnal.get('akun_kredit', ''),
            "nominal_kredit": jurnal.get('nominal_kredit', 0)
        }
        supabase.table("jurnal_akuntansi").insert(data_insert).execute()

    # Generate Excel
    excel_filename = generate_excel(data_hasil_ai, user_name)
    excel_url = f"/files/{excel_filename}"

    # Bikin pesan respons
    nama = data_hasil_ai.get('nama_perusahaan', '')
    periode = data_hasil_ai.get('periode', '')
    jumlah = len(data_hasil_ai['daftar_jurnal'])
    file_info = f"\n📎 Sumber: File {from_file}" if from_file else ""

    pesan_fun = (
        f"{ai_result.get('pesan_balasan', 'Beres nih!')} ✨{file_info}\n\n"
        f"📋 Perusahaan: {nama}\n"
        f"📅 Periode: {periode}\n"
        f"📝 Total {jumlah} entri jurnal\n"
        f"💰 Total Debit & Kredit: Rp {total_debit:,.0f}\n"
        f"{status_balance}\n\n"
        f"📥 File Excel juga udah gue siapin lengkap dengan:\n"
        f"   • Jurnal Umum\n"
        f"   • Neraca Saldo (jika ada)\n"
        f"   • Laporan Harga Pokok Produksi (jika ada)\n"
        f"   • Laporan Harga Pokok Penjualan (jika ada)\n\n"
        f"Download Excel-nya di bawah ya~ 👇"
    )

    return ChatResponse(
        type="jurnal",
        message=pesan_fun,
        data=HasilParserAkuntansi(**data_hasil_ai),
        excel_url=excel_url
    )


# ===== ENDPOINT UTAMA: CHAT PINTAR =====
@app.post("/api/chat", response_model=ChatResponse)
def chat_pintar(payload: ChatRequest):
    try:
        return _process_ai_pintar(payload.message, payload.user_name)
    except HTTPException as http_err:
        raise http_err
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ===== ENDPOINT BARU: CHAT DENGAN FILE =====
@app.post("/api/chat-with-file", response_model=ChatResponse)
async def chat_dengan_file(
    file: UploadFile = File(...),
    message: str = Form(default=""),
    user_name: str = Form(default="User"),
):
    # Validasi file
    allowed_extensions = {'.doc', '.docx', '.txt', '.pdf', '.xlsx', '.xls', '.csv'}
    ext = Path(file.filename or '').suffix.lower()
    if ext not in allowed_extensions:
        raise HTTPException(
            status_code=400,
            detail=f"Format file '{ext}' belum didukung. Coba pake: {', '.join(allowed_extensions)}"
        )

    # Simpan file sementara
    file_id = uuid.uuid4().hex[:12]
    safe_filename = f"{file_id}{ext}"
    file_path = UPLOAD_DIR / safe_filename

    try:
        with open(file_path, "wb") as buffer:
            content = await file.read()
            buffer.write(content)

        # Extract teks dari file
        extracted_text = extract_text_from_file(str(file_path), file.filename or "")

        if not extracted_text or len(extracted_text.strip()) < 10:
            raise ValueError("File sepertinya kosong atau tidak bisa dibaca. Coba file lain ya!")

        # Gabungkan dengan pesan user (kalau ada)
        full_context = f"ISI FILE '{file.filename}':\n{extracted_text}"
        if message.strip():
            full_context = f"Pesan dari user: {message}\n\n{full_context}"
            
        return _process_ai_pintar(full_context, user_name, from_file=file.filename or "file")

    except HTTPException:
        raise
    except ValueError as ve:
        raise HTTPException(status_code=400, detail=str(ve))
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        # Cleanup uploaded file
        if file_path.exists():
            file_path.unlink()


# ===== ENDPOINT DOWNLOAD EXCEL =====
@app.get("/api/download/{filename}")
def download_excel(filename: str):
    """Download file Excel yang udah di-generate"""
    filepath = EXCEL_DIR / filename
    if not filepath.exists():
        raise HTTPException(status_code=404, detail="File tidak ditemukan.")
    return FileResponse(
        path=str(filepath),
        filename=filename,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )


# ===== ENDPOINT LAMA (backward compat) =====
@app.post("/api/parse-soal", response_model=HasilParserAkuntansi)
def parse_soal_manufaktur(payload: SoalRequest):
    try:
        prompt_instruksi = f"""
        Lu adalah asisten dosen akuntansi manufaktur super pintar. Tugas utama lu merapikan soal teks 
        kronologis yang berantakan menjadi data JSON yang komprehensif, mencakup:
        1. Jurnal Umum
        2. Neraca Saldo (jika diminta atau bisa dibuat berdasarkan neraca awal + jurnal)
        3. Laporan Harga Pokok Produksi (jika ada data bahan baku, BDP, dll)
        4. Laporan Harga Pokok Penjualan (jika ada data barang jadi terjual)
        
        Aturan: COA standar, akrual ke 'Utang Beban', jurnal harus balance, hitung Neraca Saldo jika bisa.
        
        Soal:
        {payload.text}
        """

        response = generate_content_with_retry(
            model='gemini-2.5-flash',
            contents=prompt_instruksi,
            config=types.GenerateContentConfig(
                response_mime_type="application/json",
                response_schema=HasilParserAkuntansi,
                temperature=0.1
            ),
        )

        if not response or not response.text:
            raise ValueError("AI tidak mengembalikan respons teks. Coba lagi.")
        data_hasil_ai = json.loads(response.text)

        total_debit = sum(j['nominal_debit'] for j in data_hasil_ai['daftar_jurnal'])
        total_kredit = sum(j['nominal_kredit'] for j in data_hasil_ai['daftar_jurnal'])

        if total_debit != total_kredit:
            raise HTTPException(status_code=400, detail="Transaksi Gak Balance, Bro!")

        for jurnal in data_hasil_ai['daftar_jurnal']:
            data_insert = {
                "nama_perusahaan": data_hasil_ai['nama_perusahaan'],
                "periode": data_hasil_ai['periode'],
                "tanggal": jurnal['tanggal'],
                "akun_debit": jurnal['akun_debit'],
                "nominal_debit": jurnal['nominal_debit'],
                "akun_kredit": jurnal['akun_kredit'],
                "nominal_kredit": jurnal['nominal_kredit']
            }
            supabase.table("jurnal_akuntansi").insert(data_insert).execute()

        return data_hasil_ai

    except HTTPException as http_err:
        raise http_err
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/jurnal")
def ambil_semua_jurnal():
    try:
        response = supabase.table("jurnal_akuntansi").select("*").order("created_at", desc=True).execute()
        return response.data
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))